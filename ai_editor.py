from typing import Optional, List, Tuple, Dict
from dataclasses import dataclass, asdict
from pathlib import Path
import difflib
import json
import time
from datetime import datetime

from openai import OpenAI
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

from config import config
from document_processor import TafsirDocumentProcessor, TafsirBlock, BlockType


def get_system_prompt() -> str:
    return """Ты корректор (НЕ редактор!) текстов тафсира на русском языке.

⚠️ КРИТИЧЕСКИ ВАЖНО:
- Если текст БЕЗ ОШИБОК — верни ТОЛЬКО одно слово: ORIGINAL
- ЗАПРЕЩЕНО переписывать, переформулировать или "улучшать стиль"
- Разрешено ТОЛЬКО исправлять явные ошибки

РАЗРЕШЕННЫЕ ИСПРАВЛЕНИЯ (только если есть):
1. Грамматические ошибки (падежи, согласование, склонение)
2. Орфографические ошибки (опечатки, неправильное написание слов)
3. Пунктуация (пропущенные или лишние запятые, точки)

СТРОГО ЗАПРЕЩЕНО:
- Менять порядок слов
- Заменять слова синонимами
- Добавлять или удалять слова (кроме исправления явных ошибок)
- "Улучшать стиль" или делать текст "более литературным"
- Менять богословский смысл
- Менять структуру предложений

СТРОГИЕ ПРАВИЛА ТРАНСЛИТЕРАЦИИ АРАБСКИХ БУКВ:
⚠️ КРИТИЧЕСКИ ВАЖНО! НЕ МЕНЯЙ ЭТИ БУКВЫ:

- Буква 'h' (латинская) = ТОЛЬКО арабская 'ه' (ха)
  Примеры: Аллаh, Умм hани, Мухаммад ﷺ
  ЗАПРЕЩЕНО менять 'h' на 'х' (русскую)!

- Буква 'х' (русская) = арабская 'ح' (ха с точкой)
  Примеры: хадис, хадж, Рахман

- Буква 'г' = арабская 'غ' (гайн)
  Примеры: Магриб, Багдад

- Буква 'с̱' (межзубная) = арабская 'ث' (са)
  Пример: с̱икр (поминание)

- Буква 'з̱' (межзубная) = арабская 'ذ' (заль)
  Пример: з̱икр (вариант)

- Символ 'ʻ' (айн) = арабская 'ع'
  Примеры: ʻАрафат, ʻАббас

- Символ ''' (апостроф-хамза) = арабская 'ء'
  Примеры: Му'мин, Ка'ба

РЕЛИГИОЗНЫЕ ТЕРМИНЫ:
- Имя Бога: ТОЛЬКО «Аллаh» (с латинской 'h' в конце)
- Имена пророков: с заглавной буквы и уважительным тоном
- Символ ﷺ (صلى الله عليه وسلم) ОСТАВЛЯЙ БЕЗ ИЗМЕНЕНИЙ
- Арабские слова в скобках НЕ ТРОГАЙ: «الحمد» или (аль-хамд)

ОБЩИЕ ПРАВИЛА:
- НЕ меняй богословский смысл текста
- НЕ удаляй и НЕ изменяй арабские слова/фразы
- НЕ добавляй новую информацию от себя
- Сохраняй все цитаты и ссылки без изменений

ПРИМЕРЫ (Few-shot):

Input: "Во истину Аллаh велик."
Output: "Воистину Аллаh велик."
Причина: Орфографическая ошибка ("Во истину" → "Воистину")

Input: "Аллаh создал небеса и землю."
Output: ORIGINAL
Причина: Нет ошибок

Input: "Сказал Всевышний Аллаh в своей книге."
Output: ORIGINAL
Причина: Стиль простой, но грамматика верна. НЕ меняй на "Всевышний Аллаh изрек в Писании"!

Input: "Пророк Мухаммад ﷺ говорил что милосердие важно."
Output: "Пророк Мухаммад ﷺ говорил, что милосердие важно."
Причина: Пунктуация (пропущена запятая перед "что")

Input: "Этот хадис передали имам Бухари и Муслим."
Output: "Этот хадис передали имам Бухари и имам Муслим."
Причина: Грамматика (согласование: "имам" нужно повторить)

Input: "Аллах могущественен и мудр."
Output: "Аллаh могущественен и мудр."
Причина: Транслитерация ('х' → 'h' для имени Бога)

Input: "В суре Аль-Фатиха говорится о величии Творца."
Output: ORIGINAL
Причина: Нет ошибок (даже если можно сделать "красивее")

ФОРМАТ ОТВЕТА:
- Нет ошибок → верни: ORIGINAL
- Есть ошибки → верни ТОЛЬКО исправленный текст (без пояснений и комментариев)"""


def clean_ayah_text(text: str) -> str:
    text = text.strip()
    text = text.strip('﴿﴾')
    text = text.replace('«', '').replace('»', '')
    text = text.replace('"', '').replace('"', '').replace('"', '')
    text = text.replace("'", '').replace("'", '').replace("'", '')
    return text.strip()


@dataclass
class EditResult:
    block_index: int
    original_text: str
    edited_text: str
    was_changed: bool
    error: Optional[str] = None
    skipped_original: bool = False

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict) -> 'EditResult':
        return cls(**data)


class EditCache:
    def __init__(self, cache_path: str):
        self.cache_path = Path(cache_path)
        self.cache: Dict[int, EditResult] = {}
        self.metadata: dict = {}
        self._load()

    def _load(self):
        if self.cache_path.exists():
            try:
                with open(self.cache_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.metadata = data.get('metadata', {})
                    results_data = data.get('results', {})
                    for block_idx_str, result_dict in results_data.items():
                        block_idx = int(block_idx_str)
                        self.cache[block_idx] = EditResult.from_dict(result_dict)
                print(f"[CACHE] Loaded {len(self.cache)} cached results from {self.cache_path.name}")
            except Exception as e:
                print(f"[CACHE] Failed to load cache: {e}")
                self.cache = {}

    def get_result(self, block_index: int) -> Optional[EditResult]:
        return self.cache.get(block_index)

    def save_result(self, result: EditResult):
        self.cache[result.block_index] = result
        self._persist()

    def _persist(self):
        try:
            data = {
                'metadata': self.metadata,
                'results': {str(idx): result.to_dict() for idx, result in self.cache.items()}
            }
            self.cache_path.parent.mkdir(parents=True, exist_ok=True)
            with open(self.cache_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"[CACHE] Failed to save cache: {e}")

    def set_metadata(self, document_path: str, model: str, total_blocks: int):
        self.metadata = {
            'document_path': document_path,
            'model': model,
            'total_blocks': total_blocks,
            'created_at': datetime.now().isoformat(),
            'last_updated': datetime.now().isoformat()
        }

    def update_metadata(self):
        self.metadata['last_updated'] = datetime.now().isoformat()
        self.metadata['cached_blocks'] = len(self.cache)

    def clear(self):
        if self.cache_path.exists():
            self.cache_path.unlink()
        self.cache = {}
        self.metadata = {}
        print(f"[CACHE] Cache cleared: {self.cache_path.name}")


class TafsirAIEditor:
    def __init__(self):
        self.client: Optional[OpenAI] = None
        self.model = config.OPENAI_MODEL
        self._init_client()

    def _init_client(self) -> bool:
        if not config.OPENAI_API_KEY:
            print("[ERROR] OPENAI_API_KEY is not set in .env")
            return False

        try:
            self.client = OpenAI(api_key=config.OPENAI_API_KEY)
            return True
        except Exception as e:
            print(f"[ERROR] Failed to initialize OpenAI client: {e}")
            return False

    def is_ready(self) -> bool:
        return self.client is not None

    def edit_text(self, text: str, max_retries: int = 3) -> Tuple[str, Optional[str]]:
        if not self.client:
            return text, "OpenAI client not initialized"

        if not text.strip():
            return text, None

        for attempt in range(1, max_retries + 1):
            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": get_system_prompt()},
                        {"role": "user", "content": text}
                    ],
                    temperature=0.1,
                    max_tokens=len(text) * 2 + 500,
                )

                edited = response.choices[0].message.content.strip()
                return edited, None

            except Exception as e:
                error_msg = str(e)

                if attempt < max_retries:
                    wait_time = 2 ** attempt
                    print(f"[RETRY] Attempt {attempt}/{max_retries} failed: {error_msg}")
                    print(f"[RETRY] Waiting {wait_time}s before retry...")
                    time.sleep(wait_time)
                else:
                    return text, f"OpenAI API error after {max_retries} attempts: {error_msg}"

        return text, "Max retries exceeded"

    def edit_block(self, block: TafsirBlock, max_retries: int = 3) -> EditResult:
        if not block.can_process_with_ai:
            return EditResult(
                block_index=block.index,
                original_text=block.text,
                edited_text=block.text,
                was_changed=False,
                error="Block is not marked for AI processing"
            )

        edited_text, error = self.edit_text(block.text, max_retries=max_retries)

        if edited_text.strip().upper() == "ORIGINAL":
            return EditResult(
                block_index=block.index,
                original_text=block.text,
                edited_text=block.text,
                was_changed=False,
                error=None,
                skipped_original=True
            )

        was_changed = edited_text.strip() != block.text.strip()

        return EditResult(
            block_index=block.index,
            original_text=block.text,
            edited_text=edited_text,
            was_changed=was_changed,
            error=error
        )


@dataclass
class DiffOperation:
    operation: str
    text: str


class VisualDiffWriter:
    def __init__(self, source_path: str):
        self.source_path = Path(source_path)
        self.document = Document(str(source_path))

    def _compute_word_diff(self, old_text: str, new_text: str) -> List[DiffOperation]:
        old_words = old_text.split()
        new_words = new_text.split()

        matcher = difflib.SequenceMatcher(None, old_words, new_words)
        operations = []

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                text = ' '.join(old_words[i1:i2])
                operations.append(DiffOperation('equal', text))

            elif tag == 'delete':
                text = ' '.join(old_words[i1:i2])
                operations.append(DiffOperation('delete', text))

            elif tag == 'insert':
                text = ' '.join(new_words[j1:j2])
                operations.append(DiffOperation('insert', text))

            elif tag == 'replace':
                if i1 < i2:
                    old_part = ' '.join(old_words[i1:i2])
                    operations.append(DiffOperation('delete', old_part))
                if j1 < j2:
                    new_part = ' '.join(new_words[j1:j2])
                    operations.append(DiffOperation('insert', new_part))

        return operations

    def apply_ayah_brackets(self, paragraph_index: int, text: str) -> bool:
        if paragraph_index >= len(self.document.paragraphs):
            return False

        paragraph = self.document.paragraphs[paragraph_index]
        paragraph.clear()

        cleaned_text = clean_ayah_text(text)

        opening_run = paragraph.add_run("\ufd3f ")
        opening_run.font.name = "Traditional Arabic"
        opening_run.font.size = Pt(16)
        opening_run.font.color.rgb = RGBColor(139, 0, 0)
        opening_run.font.bold = False

        text_run = paragraph.add_run(cleaned_text)
        text_run.font.name = "Traditional Arabic"
        text_run.font.size = Pt(16)
        text_run.font.color.rgb = RGBColor(139, 0, 0)
        text_run.font.bold = False

        closing_run = paragraph.add_run(" \ufd3e")
        closing_run.font.name = "Traditional Arabic"
        closing_run.font.size = Pt(16)
        closing_run.font.color.rgb = RGBColor(139, 0, 0)
        closing_run.font.bold = False

        return True

    def apply_visual_diff(self, paragraph_index: int, original: str, edited: str) -> bool:
        if paragraph_index >= len(self.document.paragraphs):
            return False

        if original.strip() == edited.strip():
            return False

        paragraph = self.document.paragraphs[paragraph_index]
        paragraph.clear()

        diff_ops = self._compute_word_diff(original, edited)

        for i, op in enumerate(diff_ops):
            if not op.text:
                continue

            if i > 0 and op.operation != 'equal':
                if diff_ops[i-1].operation != 'equal':
                    paragraph.add_run(" ")

            if op.operation == 'equal':
                paragraph.add_run(op.text)

            elif op.operation == 'delete':
                run = paragraph.add_run(op.text)
                run.font.strike = True
                run.font.color.rgb = RGBColor(180, 0, 0)

            elif op.operation == 'insert':
                run = paragraph.add_run(op.text)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.color.rgb = RGBColor(0, 100, 0)

            if i < len(diff_ops) - 1:
                next_op = diff_ops[i + 1]
                if next_op.operation == 'equal':
                    paragraph.add_run(" ")

        return True

    def apply_edits(self, edit_results: List[EditResult], ayah_blocks: List[TafsirBlock] = None) -> int:
        modified_count = 0

        for result in edit_results:
            if result.was_changed and not result.error and not result.skipped_original:
                success = self.apply_visual_diff(
                    result.block_index,
                    result.original_text,
                    result.edited_text
                )
                if success:
                    modified_count += 1

        if ayah_blocks:
            for ayah in ayah_blocks:
                self.apply_ayah_brackets(ayah.index, ayah.text)

        return modified_count

    def save(self, output_path: str) -> bool:
        try:
            output = Path(output_path)
            output.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(str(output))
            print(f"[OK] Document saved: {output}")
            return True
        except Exception as e:
            print(f"[ERROR] Failed to save document: {e}")
            return False


def edit_document(
    input_path: str,
    output_path: Optional[str] = None,
    max_blocks: Optional[int] = None,
    dry_run: bool = False,
    use_cache: bool = True,
    clear_cache: bool = False
) -> Tuple[int, int, List[EditResult]]:
    if not output_path:
        input_file = Path(input_path)
        output_path = str(input_file.parent / f"{input_file.stem}_edited{input_file.suffix}")

    cache_path = f"{input_path}.cache.json"
    cache = EditCache(cache_path) if use_cache else None

    if clear_cache and cache:
        cache.clear()

    print("\n" + "=" * 70)
    print("AI-POWERED DOCUMENT CORRECTION (Surgical Word-Level Diff)")
    if use_cache:
        print("WITH RESUMABLE PROCESSING (Checkpoints)")
    print("=" * 70)
    print(f"\n  Input:  {input_path}")
    print(f"  Output: {output_path}")
    print(f"  Model:  {config.OPENAI_MODEL}")
    if use_cache:
        print(f"  Cache:  {cache_path}")
    if dry_run:
        print("  Mode:   DRY RUN (no changes will be saved)")
    print()

    editor = TafsirAIEditor()
    if not editor.is_ready():
        print("[ERROR] AI corrector not ready. Check OPENAI_API_KEY in .env")
        return 0, 0, []

    processor = TafsirDocumentProcessor()
    if not processor.load(input_path):
        return 0, 0, []

    processor.classify_document()

    ai_blocks = processor.get_ai_processable_blocks()
    ayah_blocks = processor.get_blocks_by_type(BlockType.AYAH)

    if cache and use_cache:
        cache.set_metadata(input_path, config.OPENAI_MODEL, len(ai_blocks))

    if max_blocks:
        ai_blocks = ai_blocks[:max_blocks]

    print(f"  Found {len(ai_blocks)} blocks for AI correction")
    print(f"  Found {len(ayah_blocks)} ayah blocks (will add beautiful brackets)")

    if cache and len(cache.cache) > 0:
        print(f"  [CACHE] {len(cache.cache)} blocks already cached\n")
    else:
        print()

    if not ai_blocks and not ayah_blocks:
        print("[INFO] No blocks to process")
        return 0, 0, []

    results: List[EditResult] = []
    total_changed = 0
    total_skipped = 0
    total_cached = 0

    for i, block in enumerate(ai_blocks):
        block_type = "COMMENTARY" if block.block_type == BlockType.COMMENTARY else "TRANSLATION"
        print(f"  [{i+1}/{len(ai_blocks)}] Processing {block_type} block #{block.index}...", end=" ")

        cached_result = cache.get_result(block.index) if cache else None

        if cached_result:
            result = cached_result
            print("CACHED")
            total_cached += 1
        else:
            try:
                result = editor.edit_block(block, max_retries=3)

                if cache:
                    cache.save_result(result)

                if result.error:
                    print(f"ERROR: {result.error}")
                    print("[CACHE] Progress saved. You can resume by re-running the command.")
                    break
                elif result.skipped_original:
                    print("ORIGINAL")
                    total_skipped += 1
                elif result.was_changed:
                    print("CHANGED")
                    total_changed += 1
                else:
                    print("no changes")

            except KeyboardInterrupt:
                print("\n[INTERRUPTED] Saving progress...")
                if cache:
                    cache.update_metadata()
                print("[CACHE] Progress saved. Resume by re-running the command.")
                return len(results), total_changed, results
            except Exception as e:
                print(f"FATAL ERROR: {e}")
                if cache:
                    cache.update_metadata()
                print("[CACHE] Progress saved.")
                break

        results.append(result)

        if result.was_changed and not result.skipped_original:
            total_changed += 1
        elif result.skipped_original:
            total_skipped += 1

    if cache:
        cache.update_metadata()

    print(f"\n  Processed: {len(results)}, Changed: {total_changed}, Skipped (ORIGINAL): {total_skipped}")
    if total_cached > 0:
        print(f"  [CACHE] Loaded from cache: {total_cached}")

    if not dry_run and (total_changed > 0 or ayah_blocks):
        print("\n  Applying surgical word-level diff to document...")
        writer = VisualDiffWriter(input_path)
        modified = writer.apply_edits(results, ayah_blocks)
        writer.save(output_path)
        print(f"\n  [OK] {modified} paragraphs modified with word-level diff")
        if ayah_blocks:
            print(f"  [OK] {len(ayah_blocks)} ayahs beautified with ﴿﴾ brackets (Traditional Arabic font)")

    if total_changed > 0:
        print("\n" + "-" * 70)
        print("SAMPLE CHANGES:")
        print("-" * 70)

        shown = 0
        for result in results:
            if result.was_changed and not result.skipped_original and shown < 3:
                print(f"\n  Block #{result.block_index}:")
                print(f"  OLD: {result.original_text[:100]}...")
                print(f"  NEW: {result.edited_text[:100]}...")
                shown += 1

    print("\n" + "=" * 70)

    return len(results), total_changed, results
