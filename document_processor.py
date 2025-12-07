import re
from enum import Enum
from pathlib import Path
from typing import List, Optional, Generator, Tuple
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from config import config


class BlockType(Enum):
    AYAH = "ayah"
    TRANSLATION = "translation"
    COMMENTARY = "commentary"
    EXPLANATION = "explanation"
    HEADER = "header"
    REFERENCE = "reference"
    EMPTY = "empty"
    UNKNOWN = "unknown"


@dataclass
class FontInfo:
    name: Optional[str] = None
    size: Optional[float] = None
    bold: bool = False
    italic: bool = False
    color_rgb: Optional[Tuple[int, int, int]] = None
    is_arabic_font: bool = False


@dataclass
class TafsirBlock:
    index: int
    block_type: BlockType
    text: str

    has_arabic: bool = False
    has_cyrillic: bool = False
    is_mixed: bool = False
    arabic_ratio: float = 0.0

    primary_font: Optional[str] = None
    font_size: Optional[float] = None
    is_bold: bool = False
    is_italic: bool = False
    text_color: Optional[Tuple[int, int, int]] = None
    is_red_text: bool = False

    can_process_with_ai: bool = False
    ai_processing_notes: str = ""

    word_count: int = 0
    char_count: int = 0

    _paragraph_ref: object = field(default=None, repr=False)


@dataclass
class DocumentStats:
    total_blocks: int = 0
    ayah_blocks: int = 0
    translation_blocks: int = 0
    commentary_blocks: int = 0
    explanation_blocks: int = 0
    header_blocks: int = 0
    reference_blocks: int = 0
    empty_blocks: int = 0
    unknown_blocks: int = 0

    total_words: int = 0
    total_characters: int = 0

    ai_processable_blocks: int = 0
    ai_processable_words: int = 0


class TafsirDocumentProcessor:

    ARABIC_RANGE = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]')
    CYRILLIC_RANGE = re.compile(r'[\u0400-\u04FF\u0500-\u052F]')

    ARABIC_FONTS = {
        'traditional arabic', 'arabic typesetting', 'sakkal majalla',
        'simplified arabic', 'arabic transparent', 'al-quran',
        'scheherazade', 'amiri', 'lateef', 'noto naskh arabic',
        'times new roman',
    }

    RED_THRESHOLD = 150

    def __init__(self, file_path: Optional[str] = None):
        self.file_path: Optional[Path] = Path(file_path) if file_path else None
        self.document: Optional[Document] = None
        self.blocks: List[TafsirBlock] = []
        self._stats: Optional[DocumentStats] = None

    def load(self, file_path: Optional[str] = None) -> bool:
        if file_path:
            self.file_path = Path(file_path)

        if not self.file_path:
            print("[ERROR] No file path provided")
            return False

        if not self.file_path.exists():
            print(f"[ERROR] File not found: {self.file_path}")
            return False

        if not self.file_path.suffix.lower() == '.docx':
            print(f"[ERROR] Not a .docx file: {self.file_path}")
            return False

        try:
            self.document = Document(str(self.file_path))
            self.blocks = []
            self._stats = None
            print(f"[OK] Document loaded: {self.file_path.name}")
            print(f"     Paragraphs: {len(self.document.paragraphs)}")
            return True
        except Exception as e:
            print(f"[ERROR] Failed to load document: {e}")
            return False

    def _count_script_chars(self, text: str) -> Tuple[int, int, int]:
        arabic = len(self.ARABIC_RANGE.findall(text))
        cyrillic = len(self.CYRILLIC_RANGE.findall(text))
        other = len(text) - arabic - cyrillic
        return arabic, cyrillic, other

    def _extract_font_info(self, paragraph) -> FontInfo:
        info = FontInfo()

        if not paragraph.runs:
            return info

        for run in paragraph.runs:
            if not run.text.strip():
                continue

            if run.font.name:
                info.name = run.font.name
                info.is_arabic_font = run.font.name.lower() in self.ARABIC_FONTS

            if run.font.size:
                info.size = run.font.size.pt

            info.bold = run.font.bold or False
            info.italic = run.font.italic or False

            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                info.color_rgb = (rgb[0], rgb[1], rgb[2])

            if info.name:
                break

        return info

    def _is_red_color(self, rgb: Optional[Tuple[int, int, int]]) -> bool:
        if not rgb:
            return False
        r, g, b = rgb
        return r > self.RED_THRESHOLD and g < 100 and b < 100

    def _detect_block_type(self, text: str, font_info: FontInfo,
                           arabic_ratio: float, has_arabic: bool,
                           has_cyrillic: bool, style_name: str) -> Tuple[BlockType, str]:
        text_stripped = text.strip()

        if not text_stripped:
            return BlockType.EMPTY, "Empty paragraph"

        if style_name and 'heading' in style_name.lower():
            return BlockType.HEADER, f"Style: {style_name}"

        explanation_keywords = ["объяснение:", "толкование:", "тафсир:"]
        for keyword in explanation_keywords:
            if text_stripped.lower().startswith(keyword):
                return BlockType.EXPLANATION, f"Starts with '{keyword}'"

        is_red = self._is_red_color(font_info.color_rgb)

        if arabic_ratio > 0.9 and is_red:
            return BlockType.AYAH, "Pure Arabic + red color"

        if arabic_ratio > 0.9 and font_info.is_arabic_font:
            return BlockType.AYAH, f"Pure Arabic + font: {font_info.name}"

        if arabic_ratio > 0.95:
            return BlockType.AYAH, "Pure Arabic text (>95%)"

        if arabic_ratio > 0.8 and font_info.name and 'arabic' in font_info.name.lower():
            return BlockType.AYAH, f"High Arabic ratio + {font_info.name}"

        if not has_arabic and has_cyrillic:
            if len(text_stripped) < 500:
                return BlockType.TRANSLATION, "Pure Cyrillic, moderate length"

        if has_cyrillic and arabic_ratio < 0.3:
            return BlockType.COMMENTARY, f"Cyrillic-dominant ({arabic_ratio:.0%} Arabic)"

        if has_cyrillic and has_arabic and arabic_ratio < 0.5:
            return BlockType.COMMENTARY, "Mixed text, Cyrillic-dominant"

        ref_patterns = [
            r'^\s*\[\d+\]',
            r'^\s*\(\d+:\d+\)',
            r'^\s*\d+\.\s',
            r'сура|аят|хадис',
        ]
        for pattern in ref_patterns:
            if re.search(pattern, text_stripped.lower()):
                return BlockType.REFERENCE, f"Matches reference pattern"

        if has_arabic and not has_cyrillic:
            return BlockType.AYAH, "Arabic-only (fallback)"

        return BlockType.UNKNOWN, "No matching rules"

    def classify_paragraph(self, index: int, paragraph) -> TafsirBlock:
        text = paragraph.text
        arabic_count, cyrillic_count, _ = self._count_script_chars(text)
        total_chars = len(text.replace(' ', '').replace('\n', ''))

        has_arabic = arabic_count > 0
        has_cyrillic = cyrillic_count > 0
        arabic_ratio = arabic_count / total_chars if total_chars > 0 else 0

        font_info = self._extract_font_info(paragraph)

        style_name = paragraph.style.name if paragraph.style else ""

        block_type, detection_reason = self._detect_block_type(
            text, font_info, arabic_ratio, has_arabic, has_cyrillic, style_name
        )

        can_ai = block_type in (BlockType.COMMENTARY, BlockType.TRANSLATION, BlockType.EXPLANATION)
        ai_notes = ""
        if block_type == BlockType.AYAH:
            ai_notes = "PROTECTED: Quranic text - no AI modification"
        elif can_ai:
            ai_notes = f"Can process: {detection_reason}"

        return TafsirBlock(
            index=index,
            block_type=block_type,
            text=text,
            has_arabic=has_arabic,
            has_cyrillic=has_cyrillic,
            is_mixed=has_arabic and has_cyrillic,
            arabic_ratio=arabic_ratio,
            primary_font=font_info.name,
            font_size=font_info.size,
            is_bold=font_info.bold,
            is_italic=font_info.italic,
            text_color=font_info.color_rgb,
            is_red_text=self._is_red_color(font_info.color_rgb),
            can_process_with_ai=can_ai,
            ai_processing_notes=ai_notes,
            word_count=len(text.split()) if text else 0,
            char_count=len(text),
            _paragraph_ref=paragraph
        )

    def classify_document(self) -> List[TafsirBlock]:
        if not self.document:
            raise ValueError("No document loaded. Call load() first.")

        self.blocks = []
        for i, para in enumerate(self.document.paragraphs):
            block = self.classify_paragraph(i, para)
            self.blocks.append(block)

        self._stats = None
        return self.blocks

    def get_stats(self) -> DocumentStats:
        if not self.blocks:
            self.classify_document()

        if self._stats:
            return self._stats

        stats = DocumentStats()

        for block in self.blocks:
            stats.total_blocks += 1
            stats.total_words += block.word_count
            stats.total_characters += block.char_count

            if block.block_type == BlockType.AYAH:
                stats.ayah_blocks += 1
            elif block.block_type == BlockType.TRANSLATION:
                stats.translation_blocks += 1
            elif block.block_type == BlockType.COMMENTARY:
                stats.commentary_blocks += 1
            elif block.block_type == BlockType.EXPLANATION:
                stats.explanation_blocks += 1
            elif block.block_type == BlockType.HEADER:
                stats.header_blocks += 1
            elif block.block_type == BlockType.REFERENCE:
                stats.reference_blocks += 1
            elif block.block_type == BlockType.EMPTY:
                stats.empty_blocks += 1
            else:
                stats.unknown_blocks += 1

            if block.can_process_with_ai:
                stats.ai_processable_blocks += 1
                stats.ai_processable_words += block.word_count

        self._stats = stats
        return stats

    def get_blocks_by_type(self, block_type: BlockType) -> List[TafsirBlock]:
        if not self.blocks:
            self.classify_document()
        return [b for b in self.blocks if b.block_type == block_type]

    def get_ai_processable_blocks(self) -> List[TafsirBlock]:
        if not self.blocks:
            self.classify_document()
        return [b for b in self.blocks if b.can_process_with_ai]

    def print_classification(self, limit: Optional[int] = None, show_empty: bool = False):
        if not self.blocks:
            self.classify_document()

        print(f"\n{'='*70}")
        print(f"DOCUMENT CLASSIFICATION: {self.file_path.name}")
        print(f"{'='*70}\n")

        type_icons = {
            BlockType.AYAH: "[AYAH]      ",
            BlockType.TRANSLATION: "[TRANSLATE] ",
            BlockType.COMMENTARY: "[COMMENTARY]",
            BlockType.EXPLANATION: "[EXPLAIN]   ",
            BlockType.HEADER: "[HEADER]    ",
            BlockType.REFERENCE: "[REFERENCE] ",
            BlockType.EMPTY: "[EMPTY]     ",
            BlockType.UNKNOWN: "[???]       ",
        }

        count = 0
        for block in self.blocks:
            if not show_empty and block.block_type == BlockType.EMPTY:
                continue

            if limit and count >= limit:
                print(f"\n... (showing {limit} of {len(self.blocks)} blocks)")
                break

            icon = type_icons.get(block.block_type, "[???]")
            ai_marker = " [AI-OK]" if block.can_process_with_ai else ""

            display_text = block.text[:80].replace('\n', ' ')
            if len(block.text) > 80:
                display_text += "..."

            print(f"[{block.index:4d}] {icon}{ai_marker}")
            print(f"       Arabic: {block.arabic_ratio:5.1%} | Font: {block.primary_font or 'default'}")
            print(f"       {display_text}")
            print()

            count += 1

        stats = self.get_stats()
        print(f"\n{'='*70}")
        print("CLASSIFICATION SUMMARY")
        print(f"{'='*70}")
        print(f"""
  Total blocks:     {stats.total_blocks}

  By Type:
    AYAH (Quran):     {stats.ayah_blocks:4d}  <- PROTECTED from AI
    TRANSLATION:      {stats.translation_blocks:4d}  <- Can process with AI
    COMMENTARY:       {stats.commentary_blocks:4d}  <- Can process with AI
    EXPLANATION:      {stats.explanation_blocks:4d}  <- Can process with AI
    HEADER:           {stats.header_blocks:4d}
    REFERENCE:        {stats.reference_blocks:4d}
    EMPTY:            {stats.empty_blocks:4d}
    UNKNOWN:          {stats.unknown_blocks:4d}

  AI Processing:
    Blocks for AI:    {stats.ai_processable_blocks}
    Words for AI:     {stats.ai_processable_words}
""")
        print(f"{'='*70}\n")


def create_sample_document(output_path: str = "documents/sample_tafsir.docx"):
    doc = Document()

    title = doc.add_heading('Тафсир Суры Аль-Фатиха', 0)

    doc.add_paragraph(
        'Сура «Аль-Фатиха» является первой сурой Священного Корана. '
        'Она называется также «Умм аль-Китаб» (Мать Книги) и «Ас-Сабу аль-Масани» '
        '(Семь повторяемых). Эта сура занимает особое место в исламе.'
    )

    ayah1 = doc.add_paragraph('بِسْمِ اللَّهِ الرَّحْمَٰنِ الرَّحِيمِ')
    for run in ayah1.runs:
        run.font.color.rgb = RGBColor(180, 0, 0)

    doc.add_paragraph('Во имя Аллаха, Милостивого, Милосердного.')

    doc.add_paragraph(
        'Тафсир: Эти слова являются началом всех благих дел. Мусульманин произносит '
        '«Бисмиллях» перед едой, перед чтением Корана, перед любым важным делом. '
        'Слово «Аллах» — это имя Всевышнего Господа, объединяющее все Его прекрасные имена.'
    )

    ayah2 = doc.add_paragraph('الْحَمْدُ لِلَّهِ رَبِّ الْعَالَمِينَ')
    for run in ayah2.runs:
        run.font.color.rgb = RGBColor(180, 0, 0)

    doc.add_paragraph('Хвала Аллаху, Господу миров!')

    doc.add_paragraph(
        'Слово «الحمد» (аль-хамд) означает восхваление и благодарность. '
        'Это более полное слово, чем просто «شكر» (шукр — благодарность). '
        'Господь миров — это Тот, Кто создал и поддерживает всё существующее.'
    )

    doc.add_paragraph('[Тафсир ибн Касир, том 1, стр. 25]')

    ayah3 = doc.add_paragraph('الرَّحْمَٰنِ الرَّحِيمِ')
    for run in ayah3.runs:
        run.font.color.rgb = RGBColor(180, 0, 0)

    doc.add_paragraph('Милостивого, Милосердного.')

    doc.add_paragraph(
        'Имена «Ар-Рахман» и «Ар-Рахим» оба происходят от корня «рахма» (милость). '
        'Ар-Рахман указывает на безграничную милость Аллаха ко всем созданиям, '
        'а Ар-Рахим — на особую милость к верующим в День Суда.'
    )

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    print(f"[OK] Sample document created: {output}")
    return str(output)


if __name__ == "__main__":
    import sys

    print("=" * 50)
    print("TAFSIR DOCUMENT PROCESSOR")
    print("Smart Parser with Block Classification")
    print("=" * 50)

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        print("\nNo file provided. Creating sample document...")
        file_path = create_sample_document()

    processor = TafsirDocumentProcessor()

    if processor.load(file_path):
        processor.classify_document()
        processor.print_classification(limit=30)

        ai_blocks = processor.get_ai_processable_blocks()
        print(f"\n[INFO] Found {len(ai_blocks)} blocks ready for AI processing")
